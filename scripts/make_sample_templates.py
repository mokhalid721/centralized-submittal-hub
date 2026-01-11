from pathlib import Path
from docx import Document

SAMPLE_DIR = Path("sample_templates")
SAMPLE_DIR.mkdir(parents=True, exist_ok=True)


def make_cover_letter(path: Path):
    doc = Document()
    doc.add_paragraph("«Company_Name»")
    doc.add_paragraph("«Address_Line»")
    doc.add_paragraph("«City_State_Zip_Code»")
    doc.add_paragraph("")

    doc.add_paragraph("Date: «Date_»")
    doc.add_paragraph("")
    doc.add_paragraph("To: «Project_Engineer_Contractor_Name»")
    doc.add_paragraph("Project: «Project_Name»")
    doc.add_paragraph("Contract No: «Contract_No»    Project No: «Project_Number»")
    doc.add_paragraph("")

    doc.add_paragraph("RE: Submittal «Sub_No» — «Sub_Title» (Spec Section: «Spec_Section»)")

    doc.add_paragraph("")
    doc.add_paragraph("Dear «Project_Engineer_Contractor_Name»,")

    doc.add_paragraph("")
    doc.add_paragraph("Please review the following submittal:")

    # bullet list placeholders (auto-deletes if blank)
    for i in range(1, 6):
        doc.add_paragraph(f"«BulletedInfo{i}»", style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("Disposition / Authorization Requested: «Authorization»")
    doc.add_paragraph("")

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("«Project_Manager_Name»")
    doc.add_paragraph("«Sender_Title»")

    doc.save(path)


def make_transmittal(path: Path):
    doc = Document()
    doc.add_paragraph("TRANSMITTAL", style="Title")
    doc.add_paragraph("")
    doc.add_paragraph("Project: «Project_Name»")
    doc.add_paragraph("Submittal No: «Sub_No»")
    doc.add_paragraph("Date: «Date_»")
    doc.add_paragraph("To: «Project_Engineer_Contractor_Name»")
    doc.add_paragraph("")
    doc.add_paragraph("Items transmitted:")
    for i in range(1, 6):
        doc.add_paragraph(f"«BulletedInfo{i}»", style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("Sent by: «Project_Manager_Name»")
    doc.save(path)


if __name__ == "__main__":
    cover = SAMPLE_DIR / "cover_letter_template.docx"
    trans = SAMPLE_DIR / "transmittal_template.docx"
    make_cover_letter(cover)
    make_transmittal(trans)
    print("Created:")
    print(" -", cover)
    print(" -", trans)
